import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
import os
import requests
from dotenv import load_dotenv
from datetime import datetime
import pytz
import urllib.parse

# .env 파일에서 환경 변수 로드
load_dotenv()

# API 키를 환경 변수에서 가져오기
openai_api_key = os.getenv("OPENAI_API_KEY")
google_time_zone_api_key = os.getenv("GOOGLE_TIME_ZONE_API_KEY")

st.set_page_config(
    page_title="Document NEW + EDIT + SUM + TIMEZONE",
    page_icon="📄",
)

# OpenAI 클라이언트 초기화
client = OpenAI(api_key=openai_api_key)

generation_config = {
    "temperature": 0.4,
    "top_p": 0.95,
}

# 사이드바에서 모델 선택
model_selection = st.sidebar.radio(
    "**사용할 모델을 선택하세요 :**", 
    ("Phoenix-GPT4o", "Phoenix-GPT4o-Mini"), 
    captions=("가격↑/성능↑/속도↓", "가격↓/성능↓/속도↑")
)
model_name = "gpt-4" if model_selection == "Phoenix-GPT4o" else "gpt-4o-mini"

st.title("Document NEW + EDIT + SUM + TIMEZONE")
st.caption("By Phoenix AI")

# 1. Doc-New: 생성할 문서의 키워드를 입력받는 부분
st.header("1. Doc-New")
keyword = st.text_input("생성할 문서의 키워드를 입력해 주세요:")

st.caption("생성한 문서의 출력 언어를 선택하세요")
output_language_new = st.selectbox(
    "",
    ("한국어", "영어", "일본어", "중국어", "러시아어", "프랑스어", "독일어", "이탈리아어")
)
language_prompts = {
    "한국어": "이 키워드에 대한 2,000자 길이의 문서를 한국어로 생성해줘.",
    "영어": "Generate a 2,000-character document for this keyword in English.",
    "일본어": "このキーワードについて2,000文字の日本語のドキュメントを作成してください。",
    "중국어": "请用中文生成关于这个关键词的2,000字文档。",
    "러시아어": "Создайте документ на 2,000 символов по этому ключевому слову на русском языке.",
    "프랑스어": "Générez un document de 2,000 caractères pour ce mot-clé en français.",
    "독일어": "Erstellen Sie ein 2,000 Zeichen langes Dokument für dieses Schlüsselwort auf Deutsch.",
    "이탈리아어": "Genera un documento di 2,000 caratteri per questa parola chiave in italiano."
}
generate_document = st.button("문서 생성")

if generate_document and keyword:
    with st.spinner("문서를 생성하는 중입니다..."):
        system_instruction = language_prompts[output_language_new]
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": keyword}
                ],
                max_tokens=2000,
                temperature=generation_config["temperature"],
                top_p=generation_config["top_p"]
            )
            
            if 'result_text' not in st.session_state:
                st.session_state.result_text = ""
            result_text = st.empty()
            result_text.success(response.choices[0].message.content.strip())
            st.session_state.result_text = response.choices[0].message.content.strip()
            with st.expander("📋 마크다운 복사"):
                st.code(st.session_state.result_text, language='markdown')
        except Exception as e:
            st.error(f"오류가 발생했습니다: {str(e)}")

    # MS Word 문서 생성 및 다운로드 기능 추가
    if 'result_text' in st.session_state and st.session_state.result_text:
        document = Document()
        document.add_heading('Generated Document', level=1)
        document.add_paragraph(st.session_state.result_text)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="문서 다운로드 (MS Word)",
            data=buffer,
            file_name=f"{keyword}_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# 2. Doc-Edit: 수정할 문서나 링크를 업로드하는 부분
st.header("2. Doc-Edit")
uploaded_file_edit = st.file_uploader("수정할 문서를 업로드 해 주세요", type=["docx"], key="edit_file")
uploaded_link_edit = st.text_input("수정할 문서의 링크를 입력해 주세요:", key="edit_link")

# 변수 초기화
doc_text_edit = ""

if uploaded_file_edit:
    document = Document(uploaded_file_edit)
    doc_text_edit = "\n".join([para.text for para in document.paragraphs])
    st.header("수정할 문서 내용")
    st.text_area("문서 내용", doc_text_edit, height=300)
elif uploaded_link_edit:
    try:
        response = requests.get(uploaded_link_edit)
        if response.status_code == 200:
            doc_text_edit = response.text
            st.header("수정할 문서 내용")
            st.text_area("문서 내용", doc_text_edit, height=300)
        else:
            st.error("문서 링크를 불러오는 데 실패했습니다.")
    except Exception as e:
        st.error(f"문서 링크를 불러오는 도중 에러가 발생했습니다: {str(e)}")

if doc_text_edit:
    edit_keyword = st.text_input("수정할 키워드 또는 문장을 입력해 주세요:")
    st.header("수정한 문서의 출력 언어를 선택하세요")
    output_language_edit = st.selectbox(
        "수정한 문서의 출력 언어를 선택하세요:",
        ("한국어", "영어", "일본어", "중국어", "러시아어", "프랑스어", "독일어", "이탈리아어")
    )
    # 여기에 문서 수정 로직을 추가할 수 있습니다.

# 3. Doc-Sum: 문서 요약 기능
st.header("3. Doc-Sum")
uploaded_file_sum = st.file_uploader("요약할 문서를 업로드 해 주세요", type=["docx"], key="sum_file")

if uploaded_file_sum:
    document = Document(uploaded_file_sum)
    doc_text_sum = "\n".join([para.text for para in document.paragraphs])
    st.header("문서 내용 요약")
    
    if st.button("요약 생성"):
        with st.spinner("요약을 생성하는 중입니다..."):
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": "Summarize the following document."},
                        {"role": "user", "content": doc_text_sum}
                    ],
                    max_tokens=500,
                    temperature=generation_config["temperature"],
                    top_p=generation_config["top_p"]
                )
                
                st.success(response.choices[0].message.content.strip())
            except Exception as e:
                st.error(f"오류가 발생했습니다: {str(e)}")

# 4. Timezone: 사용자가 입력한 장소의 시간 확인
st.header("4. Timezone")
location = st.text_input("시간을 확인할 장소를 입력하세요:")

if st.button("시간 확인") and location:
    try:
        # 주소를 URL 인코딩
        encoded_location = urllib.parse.quote(location)
        geocode_url = f"https://maps.googleapis.com/maps/api/geocode/json?address={encoded_location}&key={google_time_zone_api_key}"
        geocode_response = requests.get(geocode_url).json()

        if geocode_response["status"] == "OK":
            latlng = geocode_response["results"][0]["geometry"]["location"]
            timezone_url = f"https://maps.googleapis.com/maps/api/timezone/json?location={latlng['lat']},{latlng['lng']}&timestamp={int(datetime.now().timestamp())}&key={google_time_zone_api_key}"
            timezone_response = requests.get(timezone_url).json()
            if timezone_response["status"] == "OK":
                tz_id = timezone_response["timeZoneId"]
                tz_name = timezone_response["timeZoneName"]
                local_time = datetime.now(pytz.timezone(tz_id))
                st.success(f"{location}의 현재 시간: {local_time.strftime('%Y-%m-%d %H:%M:%S')} ({tz_name})")
            else:
                st.error("시간대를 불러오는 데 실패했습니다.")
        else:
            st.error(f"위치를 찾을 수 없습니다. 상태 코드: {geocode_response['status']}, 메시지: {geocode_response.get('error_message', '없음')}")
    except Exception as e:
        st.error(f"시간을 확인하는 도중 에러가 발생했습니다: {str(e)}")
